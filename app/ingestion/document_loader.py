"""
Document Loader - Handles loading various document formats
Supports: PDF, DOCX, TXT, Markdown
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
import structlog

logger = structlog.get_logger()


@dataclass
class Document:
    """Represents a loaded document with metadata"""
    content: str
    source: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __post_init__(self):
        if not self.metadata:
            self.metadata = {}
        self.metadata["source"] = self.source


class DocumentLoader:
    """
    Loads documents from various file formats.
    Extensible design allows adding new format handlers.
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}
    
    def __init__(self, documents_path: str = "data/documents"):
        self.documents_path = Path(documents_path)
        self.documents_path.mkdir(parents=True, exist_ok=True)
        
    def load_all(self) -> List[Document]:
        """Load all documents from the documents directory"""
        documents = []
        
        for file_path in self.documents_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(file_path)
                    if doc:
                        documents.append(doc)
                        logger.info("Loaded document", file=str(file_path))
                except Exception as e:
                    logger.error("Failed to load document", file=str(file_path), error=str(e))
        
        logger.info("Document loading complete", total_documents=len(documents))
        return documents
    
    def load_file(self, file_path: Path) -> Optional[Document]:
        """Load a single file based on its extension"""
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            return self._load_pdf(file_path)
        elif extension == ".docx":
            return self._load_docx(file_path)
        elif extension in {".txt", ".md", ".markdown"}:
            return self._load_text(file_path)
        else:
            logger.warning("Unsupported file format", file=str(file_path))
            return None
    
    def _load_pdf(self, file_path: Path) -> Optional[Document]:
        """Load PDF document using pypdf"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_type": "pdf",
                    "page_count": len(reader.pages),
                    "filename": file_path.name
                }
            )
        except ImportError:
            logger.error("pypdf not installed. Run: pip install pypdf")
            return None
        except Exception as e:
            logger.error("Error loading PDF", file=str(file_path), error=str(e))
            return None
    
    def _load_docx(self, file_path: Path) -> Optional[Document]:
        """Load DOCX document using python-docx"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)
            
            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_type": "docx",
                    "paragraph_count": len(paragraphs),
                    "filename": file_path.name
                }
            )
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None
        except Exception as e:
            logger.error("Error loading DOCX", file=str(file_path), error=str(e))
            return None
    
    def _load_text(self, file_path: Path) -> Optional[Document]:
        """Load plain text or markdown file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            return Document(
                content=content,
                source=str(file_path),
                metadata={
                    "file_type": file_path.suffix.lower().strip("."),
                    "filename": file_path.name
                }
            )
        except Exception as e:
            logger.error("Error loading text file", file=str(file_path), error=str(e))
            return None
    
    def add_document_from_text(self, text: str, source_name: str) -> Document:
        """Create a document from raw text (useful for API integrations)"""
        return Document(
            content=text,
            source=source_name,
            metadata={
                "file_type": "text",
                "filename": source_name
            }
        )
